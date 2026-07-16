#!/usr/bin/env python3
"""Create SmartNPS360 Simple Tester Guide as a Word doc (open in Google Docs)."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "SmartNPS360_Simple_Tester_Guide.docx"
NAVY = RGBColor(0x02, 0x2A, 0x67)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)
NOTE_HEX = "FFFBF0"
ALT_HEX = "F5F8FC"
HEADER_HEX = "022A67"


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, bold=False, size=10, color=BODY):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Arial"


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, bold=True, size=12, color=NAVY)


def add_body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=9, color=MUTED if italic else BODY)
    r.italic = italic


def add_checklist_table(doc, rows):
    """rows: list of (what, expected)"""
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Done", "What to test", "Expected result", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, bold=True, size=9, color=WHITE)
        set_cell_shading(cell, HEADER_HEX)

    for r_idx, (what, expected) in enumerate(rows, start=1):
        values = ["☐", what, expected, ""]
        for c_idx, value in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run(r, size=9, color=BODY)
            if c_idx == 3:
                set_cell_shading(cell, NOTE_HEX)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, ALT_HEX)

    # Approximate column widths
    widths = [Cm(1.4), Cm(5.4), Cm(6.2), Cm(3.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    doc.add_paragraph()


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SmartNPS360")
    set_run(r, bold=True, size=20, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("App Testing Checklist")
    set_run(r, size=12, color=MUTED)

    add_body(
        doc,
        "Use this list while testing the SmartNPS360 app on your phone. "
        "Tick Done when finished. In the yellow Notes column write Pass or Fail. "
        "If something fails, write what you saw.",
    )

    add_heading(doc, "1. Open the app & internet")
    add_checklist_table(
        doc,
        [
            (
                "Open SmartNPS360 from your phone",
                "App opens and shows the SmartNPS360 website (login or home).",
            ),
            (
                "Turn on Airplane Mode, then open or reload the app",
                "You see “No internet connection” and a Retry button.",
            ),
            (
                "Turn internet back on and tap Retry",
                "The page loads again.",
            ),
        ],
    )

    add_heading(doc, "2. Login & logout")
    add_checklist_table(
        doc,
        [
            (
                "On the login page, enter a wrong password and submit",
                "Login fails. You stay on the login page.",
            ),
            (
                "Log in with a correct account",
                "You enter the officer area of the app.",
            ),
            (
                "Fully close the app, then open it again (while still logged in)",
                "You can continue without problems (session still works).",
            ),
            (
                "While your shift has ended (off duty): log out",
                "You return to the login page.",
            ),
            (
                "While still on shift (on duty): try to log out",
                "Logout is blocked. Message says you must end your shift first.",
            ),
        ],
    )

    add_heading(doc, "3. Bottom menu (after login)")
    add_body(
        doc,
        "The bottom menu shows only on the main Dashboard, TimeSheet, and Profile pages.",
        italic=True,
    )
    add_checklist_table(
        doc,
        [
            ("Tap Dashboard", "Dashboard page opens. Menu item looks selected."),
            ("Tap TimeSheet", "Monthly timesheet page opens."),
            ("Tap Profile", "Profile page opens."),
        ],
    )

    add_heading(doc, "4. Clock in / shift attendance")
    add_body(
        doc,
        "Phone Location must be on. The app may show “Location required for shift attendance” "
        "with buttons Cancel and Continue. You may need location set to Always / Allow all the time.",
        italic=True,
    )
    add_checklist_table(
        doc,
        [
            (
                "Start clock-in / shift attendance from the app",
                "If a location message appears, title is about location for shift attendance.",
            ),
            (
                "On that message, tap Continue (then allow location when the phone asks)",
                "You can move to the next location step.",
            ),
            (
                "If asked, set location to Always / Allow all the time",
                "Clock-in can complete.",
            ),
            (
                "Clock in outdoors or near a window (good GPS)",
                "Attendance / clock-in succeeds.",
            ),
            (
                "On the location message, tap Cancel",
                "Clock-in stops. You are not forced to continue.",
            ),
            (
                "Turn Location off on the phone, then try clock-in",
                "You may see “Turn on location services” or Open Settings / Cancel.",
            ),
            (
                "If Open Settings is shown: open Settings, fix location, return to the app, try again",
                "Clock-in works after location is fixed.",
            ),
            (
                "If the phone uses fake/mock GPS, try clock-in",
                "You see “Mock location detected” and an OK button. Clock-in does not succeed until mock GPS is off.",
            ),
        ],
    )

    add_heading(doc, "5. During your shift (on duty)")
    add_checklist_table(
        doc,
        [
            (
                "Start / stay on duty with location fully allowed",
                "Shift stays active. On Android you may see a notification: “Sharing live location”.",
            ),
            (
                "Leave the app for a few minutes while still on duty (open another app)",
                "App does not crash. You can return and continue.",
            ),
            (
                "If a top banner asks for location (button like Enable Location / Allow Location)",
                "Tapping the button helps you fix location settings.",
            ),
            (
                "End your shift (go off duty)",
                "Location sharing stops. You can log out after this.",
            ),
        ],
    )

    add_heading(doc, "6. Camera & photos")
    add_body(
        doc,
        "These appear when a page asks for a photo (for example reports, visits, or profile / ID).",
        italic=True,
    )
    add_checklist_table(
        doc,
        [
            (
                "When asked for a photo, allow Camera and take a picture",
                "Camera opens. Photo can be attached.",
            ),
            (
                "Cancel without taking a photo",
                "You return safely. App still works.",
            ),
            (
                "Choose a photo from your gallery when asked",
                "You can pick an image and attach it.",
            ),
            (
                "Deny Camera permission, then try to take a photo again",
                "Photo does not work. App shows that permission was denied.",
            ),
        ],
    )

    add_heading(doc, "7. Notifications")
    add_checklist_table(
        doc,
        [
            (
                "Allow notifications when the phone or app asks",
                "Notifications are allowed.",
            ),
            (
                "Ask your team to send a test alert while the app is open",
                "You see the notification (sound if the phone is not silent).",
            ),
            (
                "Receive an alert with the app closed, then tap the notification",
                "App opens to a SmartNPS360 page (often Dashboard).",
            ),
            (
                "Turn off notifications for SmartNPS360 in phone Settings",
                "Alerts stop until you turn them on again.",
            ),
        ],
    )

    add_heading(doc, "8. Other actions")
    add_checklist_table(
        doc,
        [
            (
                "On Profile, update your details and save (if the page allows)",
                "Changes are saved.",
            ),
            (
                "If a page offers Share, tap it",
                "Phone share options open.",
            ),
            (
                "If a page offers Download, download a file",
                "Download completes. App does not crash.",
            ),
            (
                "If the website has light / dark mode, switch it",
                "App look updates to match.",
            ),
        ],
    )

    add_heading(doc, "9. Full use cases")
    add_body(
        doc,
        "Do these end-to-end. For clock-in, test both success and every way clock-in can fail "
        "(location off, permission denied, Always missing, Cancel, mock GPS, poor GPS). "
        "After each failure, fix the setting and confirm clock-in works again.",
        italic=True,
    )
    add_checklist_table(
        doc,
        [
            (
                "Open app → log in → tap Dashboard, TimeSheet, and Profile",
                "All three pages open. No freeze.",
            ),
            (
                "Wrong password, then correct password",
                "Fails first, then successful login.",
            ),
            (
                "Get a notification → tap it",
                "App opens to the related SmartNPS360 page.",
            ),
            (
                "On a page that needs a photo: allow Camera → take photo → attach → save / submit",
                "Photo is attached and saved.",
            ),
            (
                "On a photo page: deny Camera → try again → allow Camera → take photo",
                "Blocked while denied; works after Camera is allowed.",
            ),
            (
                "SUCCESS: Log in → clock in → on location message tap Continue → allow location → set Always / Allow all the time → finish clock-in outdoors",
                "Clock-in succeeds. You are on shift.",
            ),
            (
                "SUCCESS: Stay on duty a few minutes → leave the app briefly → return → end shift → log out",
                "Full day path works. Logout works only after the shift ends.",
            ),
            (
                "FAIL: Turn Location / Location Services OFF on the phone → try clock-in",
                "Unable to clock in. You see “Turn on location services” (or Open Settings / Cancel). Clock-in does not complete.",
            ),
            (
                "RECOVER: Turn Location back ON → try clock-in again (with Always allowed)",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: On “Location required for shift attendance”, tap Cancel",
                "Unable to clock in. Flow stops. You are not forced to continue.",
            ),
            (
                "RECOVER: Start clock-in again → tap Continue → allow what is needed → finish",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: When the phone asks for location, tap Don’t Allow / Deny",
                "Unable to clock in. App asks you to open Settings or shows location is needed.",
            ),
            (
                "RECOVER: Open Settings from the app (or phone Settings) → allow location → return → clock in again",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: Allow location only While Using / only this time (not Always / All the time) → try clock-in",
                "Unable to clock in. App asks for background / Always location (Open Settings / Cancel).",
            ),
            (
                "RECOVER: In Settings set SmartNPS360 location to Always / Allow all the time → return → clock in",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: On Open Settings message, tap Cancel (do not open Settings)",
                "Unable to clock in. Attendance is not completed.",
            ),
            (
                "FAIL: Tap Open Settings but leave location wrong → return to the app",
                "Unable to clock in. You may see “Unable to verify attendance”.",
            ),
            (
                "RECOVER: Set location correctly to Always → return → clock in again",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: Turn on fake / mock GPS on the phone → try clock-in (with Always already allowed)",
                "Unable to clock in. You see “Mock location detected” and OK.",
            ),
            (
                "RECOVER: Turn mock GPS off → try clock-in again",
                "Clock-in succeeds.",
            ),
            (
                "FAIL: Try clock-in indoors with very poor GPS / weak signal (Always already allowed)",
                "Clock-in may fail or take long if GPS is not accurate enough. Retry outdoors or near a window.",
            ),
            (
                "RECOVER: Move outdoors / near a window → clock in again",
                "Clock-in succeeds with a good GPS fix.",
            ),
            (
                "FAIL: Double-tap clock-in quickly while the first permission step is still running",
                "Second try may say to wait until the location step finishes. First flow continues.",
            ),
            (
                "COMBINED: Location OFF + no Always permission → try clock-in → fix Location ON → still only While Using → try again → set Always → clock in",
                "Blocked at each wrong step. Succeeds only when Location is on and Always is set.",
            ),
        ],
    )

    add_heading(doc, "Sign-off")
    sign = doc.add_table(rows=3, cols=4)
    sign.style = "Table Grid"
    labels = [
        ("Your name", "", "Phone", ""),
        ("Date", "", "Overall (Pass / Fail)", ""),
        ("Problems found", "", "", ""),
    ]
    for r_idx, row_vals in enumerate(labels):
        for c_idx, val in enumerate(row_vals):
            cell = sign.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_run(r, bold=bool(val and c_idx in (0, 2)), size=9, color=NAVY if val else BODY)
            if c_idx in (0, 2) and val:
                set_cell_shading(cell, "EEF2F7")
            elif c_idx in (1, 3) or (r_idx == 2 and c_idx >= 1):
                set_cell_shading(cell, NOTE_HEX)

    # Merge problems found value cells
    sign.rows[2].cells[1].merge(sign.rows[2].cells[3])

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
